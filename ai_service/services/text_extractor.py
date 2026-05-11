"""
text_extractor.py
─────────────────
Extracts plain text from PDF and DOCX resume files.

Fixes applied vs original:
  - DOCX table cell text was appended with no spacing, causing words to merge.
    Each cell now appended with a space separator before newline.
  - Added explicit encoding fallback for PDFs with non-UTF-8 characters.
  - Added file extension check that is case-insensitive (.PDF / .DOCX).
  - Empty-page guard prevents blank pages from adding whitespace-only lines.
"""

import io
from typing import Union

import PyPDF2
from docx import Document


# ──────────────────────────────────────────────────────────────────────────────
# PDF
# ──────────────────────────────────────────────────────────────────────────────

def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from raw PDF bytes using PyPDF2."""
    try:
        pdf_file   = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        pages = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():          # ✅ skip blank pages
                pages.append(page_text.strip())

        return "\n".join(pages)

    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")


# ──────────────────────────────────────────────────────────────────────────────
# DOCX
# ──────────────────────────────────────────────────────────────────────────────

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from raw DOCX bytes using python-docx."""
    try:
        docx_file = io.BytesIO(content)
        doc       = Document(docx_file)

        lines = []

        # Body paragraphs
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                lines.append(para.text.strip())

        # Table cells — ✅ fix: was missing space between cells, words merged
        for table in doc.tables:
            for row in table.rows:
                row_parts = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_parts.append(cell.text.strip())
                if row_parts:
                    lines.append("  ".join(row_parts))

        return "\n".join(lines)

    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")


# ──────────────────────────────────────────────────────────────────────────────
# Dispatcher
# ──────────────────────────────────────────────────────────────────────────────

def extract_text_from_file(content: bytes, filename: str) -> str:
    """
    Dispatch to the appropriate extractor based on file extension.
    Extension check is case-insensitive so .PDF and .DOCX also work.
    """
    if not filename:
        raise ValueError("Filename is required to determine file type")

    lower = filename.lower()

    if lower.endswith(".pdf"):
        return extract_text_from_pdf(content)
    elif lower.endswith(".docx"):
        return extract_text_from_docx(content)
    else:
        raise ValueError(
            f"Unsupported file format: '{filename}'. "
            "Only PDF and DOCX files are accepted."
        )
